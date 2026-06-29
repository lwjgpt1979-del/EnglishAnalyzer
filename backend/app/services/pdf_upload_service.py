"""教材 PDF 上传 + 文本提取 + 单元自动检测（M3）。

流程：
  1. save_upload()       — 保存 bytes 到 uploads/pdfs/<file_id>.pdf
  2. extract_pages()     — pdfplumber 逐页提取文本，返回 list[str]（index 0 = 第 1 页）
  3. auto_detect_units() — 扫描页文本识别单元分界，返回 list[dict] | None
  4. get_unit_text()     — 拼接指定页范围文本，供 AI 生成使用
"""
from __future__ import annotations

import logging
import re
import uuid
from pathlib import Path

_log = logging.getLogger(__name__)

UPLOAD_DIR = Path(__file__).parent.parent.parent / "uploads" / "pdfs"
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

# 常见单元标题匹配模式
_UNIT_PATTERNS = [
    re.compile(r"\bUnit\s+(\d+)\b", re.IGNORECASE),
    re.compile(r"\bUnit\s+(One|Two|Three|Four|Five|Six|Seven|Eight|Nine|Ten)\b", re.IGNORECASE),
    re.compile(r"第\s*([一二三四五六七八九十\d]+)\s*单元"),
    re.compile(r"\bStarter\s+Unit\b", re.IGNORECASE),
]

_WORD_TO_INT: dict[str, int] = {
    "one": 1, "two": 2, "three": 3, "four": 4, "five": 5,
    "six": 6, "seven": 7, "eight": 8, "nine": 9, "ten": 10,
}
_CN_TO_INT: dict[str, int] = {
    "一": 1, "二": 2, "三": 3, "四": 4, "五": 5,
    "六": 6, "七": 7, "八": 8, "九": 9, "十": 10,
}


# ─── 文件存储 ─────────────────────────────────────────────────────────────────

def save_upload(file_bytes: bytes) -> str:
    """保存上传文件，返回 file_id（hex UUID，不含扩展名）。"""
    file_id = uuid.uuid4().hex
    (UPLOAD_DIR / f"{file_id}.pdf").write_bytes(file_bytes)
    return file_id


def delete_upload(file_id: str) -> None:
    """生成完成后清理临时文件（optional）。"""
    (UPLOAD_DIR / f"{file_id}.pdf").unlink(missing_ok=True)


def save_upload_docx(file_bytes: bytes) -> str:
    """保存上传的 Word(.docx)，返回 file_id（hex UUID，不含扩展名）。"""
    file_id = uuid.uuid4().hex
    (UPLOAD_DIR / f"{file_id}.docx").write_bytes(file_bytes)
    return file_id


def extract_docx_text(file_id: str) -> str:
    """提取 .docx 全文，**按文档真实顺序**交错段落与表格。需要 python-docx。

    关键：python-docx 的 doc.paragraphs / doc.tables 会把所有表格甩到段落之后，
    导致卷中内联的阅读框（如 Noticeboard）、选词框与其题目脱节。这里改为遍历
    body 子节点，按出现次序还原，保证拆题时材料挂回正确的题。
    """
    try:
        from docx import Document
        from docx.oxml.ns import qn
        from docx.table import Table
        from docx.text.paragraph import Paragraph
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError("python-docx 未安装，请 pip install python-docx") from exc

    path = UPLOAD_DIR / f"{file_id}.docx"
    if not path.exists():
        raise FileNotFoundError(f"DOCX not found: {file_id}")

    doc = Document(str(path))
    parts: list[str] = []
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            text = Paragraph(child, doc).text
            if text.strip():
                parts.append(text)
        elif child.tag == qn("w:tbl"):
            for row in Table(child, doc).rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append("\t".join(cells))
    return "\n".join(parts).strip()


# ─── 文本提取 ─────────────────────────────────────────────────────────────────

def _ocr_sidecar(file_id: str) -> Path:
    return UPLOAD_DIR / f"{file_id}.ocr.json"


def ocr_text_available(file_id: str) -> bool:
    return _ocr_sidecar(file_id).exists()


def extract_pages(file_id: str) -> list[str]:
    """逐页提取文本,返回 list(index 0 = 第 1 页)。

    若存在 OCR 旁路文件(扫描件已 OCR)→ 直接用 OCR 文字;否则 pdfplumber 抽文字层。
    这样单元检测与生成(get_unit_text)对"扫描件已 OCR"完全透明。
    """
    import json
    side = _ocr_sidecar(file_id)
    if side.exists():
        try:
            return list(json.loads(side.read_text(encoding="utf-8")))
        except Exception:  # noqa: BLE001
            pass
    try:
        import pdfplumber
    except ImportError as exc:
        raise RuntimeError("pdfplumber 未安装，请 pip install pdfplumber") from exc

    path = UPLOAD_DIR / f"{file_id}.pdf"
    if not path.exists():
        raise FileNotFoundError(f"PDF not found: {file_id}")

    pages: list[str] = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            pages.append(text.strip())
    return pages


def extract_text_from_pdf_bytes(pdf_bytes: bytes) -> str:
    """从 PDF bytes 逐页抽文字层并拼接。用于从已拆出的单元 PDF(unit_pdf_url)回取原文。

    扫描件(无文字层)会返回空串——调用方据此提示走 OCR 或重传文字版。需 pdfplumber。
    """
    import io
    try:
        import pdfplumber
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError("pdfplumber 未安装，请 pip install pdfplumber") from exc
    parts: list[str] = []
    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
        for page in pdf.pages:
            text = (page.extract_text() or "").strip()
            if text:
                parts.append(text)
    return "\n\n---\n\n".join(parts)


async def ocr_pdf_bytes(pdf_bytes: bytes, *, resolution: int = 130,
                        concurrency: int = 6, on_progress=None) -> str:
    """扫描件 PDF bytes(无文字层)→ 逐页渲染图 → 豆包视觉 OCR → 拼接原样文字。

    与 ocr_pages_to_sidecar 同机制,但作用于 bytes 且直接返回文字(不落 sidecar),
    供「单元 PDF(unit_pdf_url)是扫描件」时按需回取原文。doubao dev 模式返回空。
    """
    import asyncio
    import base64
    import io

    import pdfplumber
    from app.services.doubao_vision_service import recognize_page_text

    def _render_all() -> list[str]:
        urls: list[str] = []
        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            for pg in pdf.pages:
                pil = pg.to_image(resolution=resolution).original.convert("RGB")
                buf = io.BytesIO()
                pil.save(buf, format="JPEG", quality=80)
                urls.append("data:image/jpeg;base64," + base64.b64encode(buf.getvalue()).decode())
        return urls

    urls = await asyncio.to_thread(_render_all)
    n = len(urls)
    results = [""] * n
    done = 0
    sem = asyncio.Semaphore(concurrency)

    async def _one(i: int) -> None:
        nonlocal done
        async with sem:
            results[i] = (await recognize_page_text(urls[i]) or "").strip()
            done += 1
            if on_progress:
                on_progress(done, n)

    await asyncio.gather(*(_one(i) for i in range(n)))
    return "\n\n---\n\n".join(r for r in results if r)


async def fetch_pdf_text(url: str, *, ocr_fallback: bool = False) -> str:
    """下载远程 PDF(如单元 unit_pdf_url 的 COS 直链)并抽文字层,失败抛异常。

    ocr_fallback=True 且文字层为空(扫描件)→ 自动逐页渲染走豆包 OCR 回取文字。
    """
    import httpx
    async with httpx.AsyncClient(timeout=60.0, follow_redirects=True) as client:
        resp = await client.get(url)
        resp.raise_for_status()
        pdf_bytes = resp.content
    text = extract_text_from_pdf_bytes(pdf_bytes)
    if not text and ocr_fallback:
        text = await ocr_pdf_bytes(pdf_bytes)
    return text


# ─── 单元自动检测 ─────────────────────────────────────────────────────────────

def _parse_unit_no(m: re.Match) -> int | None:
    """从正则匹配提取单元编号（整数）；Starter Unit → None（调用方处理）。"""
    if not m.lastindex:
        return None  # Starter Unit 等无编号模式
    raw = m.group(1)
    if isinstance(raw, str) and raw.isdigit():
        return int(raw)
    raw_lower = raw.lower() if isinstance(raw, str) else ""
    if raw_lower in _WORD_TO_INT:
        return _WORD_TO_INT[raw_lower]
    if raw in _CN_TO_INT:
        return _CN_TO_INT[raw]
    try:
        return int(raw)
    except (ValueError, TypeError):
        return None


# 末单元尾部常跟 Workbook / 词汇表 / 不规则动词表 / 附录 等非单元正文 → 据此裁剪末单元结束页。
_BACKMATTER_PATTERNS = [
    re.compile(r"\bWorkbook\b", re.I),
    re.compile(r"\bWord\s?list\b", re.I),
    re.compile(r"\bVocabulary\b", re.I),
    re.compile(r"\bIrregular\s+verbs\b", re.I),
    re.compile(r"\bGrammar\s+reference\b", re.I),
    re.compile(r"\bAppendix\b", re.I),
    re.compile(r"词汇表|不规则动词|附\s*录"),
]


def _backmatter_start(pages: list[str], after_page: int) -> int | None:
    """从 after_page(1-based,不含)之后找首个后置附录页(Workbook/词汇表/附录…),返回其 1-based 页码。"""
    for idx in range(after_page, len(pages)):   # idx 0-based → 1-based 页 = idx+1 > after_page
        head = (pages[idx] or "")[:120]
        if any(pat.search(head) for pat in _BACKMATTER_PATTERNS):
            return idx + 1
    return None


def detect_page_offset(pages: list[str]) -> int:
    """估算 PDF 页序与书本印刷页码的固定偏移(印刷页码 = PDF 页序 − offset)。

    每页文本几乎都含自己的印刷页码,用"PDF 页序 − 页内数字"投票:真实偏移会被
    几乎每页投到、拿到压倒性多数;练习号/年份等其它数字散票。不显著则返回 0(按 PDF 页序显示)。
    """
    from collections import Counter
    votes: Counter = Counter()
    for i, text in enumerate(pages, start=1):
        for m in re.findall(r"\b(\d{1,3})\b", text or ""):
            k = i - int(m)
            if 0 <= k <= 30:               # 前置页(封面/版权/目录)数量合理范围
                votes[k] += 1
    if not votes:
        return 0
    top = votes.most_common(2)
    best, cnt = top[0]
    second = top[1][1] if len(top) > 1 else 0
    # 需足够多页支持且明显领先次高,否则视为不确定
    if cnt >= max(5, len(pages) // 3) and cnt >= second * 2:
        return best
    return 0


def auto_detect_units(pages: list[str]) -> list[dict] | None:
    """
    扫描各页文本，识别单元起始页。

    返回 list[dict(unit_no, start_page, end_page, detected_title)]，
    页码均为 1-based；若识别单元数 < 2 则返回 None（触发人工辅助流程）。

    检测策略：只看每页前 500 字符，避免被正文内容的 "Unit" 字样误触发。
    """
    hits: list[tuple[int, int, str]] = []  # (page_1based, unit_no, title_snippet)

    for idx, text in enumerate(pages):
        page_no = idx + 1
        first_line = (text.split("\n")[0][:80] if text else "").strip()
        head = text[:500]
        for pat in _UNIT_PATTERNS:
            m = pat.search(head)
            if m:
                unit_no = _parse_unit_no(m)
                if unit_no is None:
                    unit_no = 0  # Starter Unit → 暂存为 0，后续重排
                # 同一单元只保留首次出现的页
                if not any(h[1] == unit_no for h in hits):
                    hits.append((page_no, unit_no, first_line))
                break

    if len(hits) < 2:
        return None

    hits.sort(key=lambda h: h[0])
    total = len(pages)

    # 重排 unit_no，保证连续（处理 Starter Unit = 0 的情况）
    for i, (pg, uno, title) in enumerate(hits):
        if uno == 0:
            hits[i] = (pg, i + 1, title)

    segments = []
    for i, (start, unit_no, title) in enumerate(hits):
        if i + 1 < len(hits):
            end = hits[i + 1][0] - 1
        else:
            # 末单元:没有下一单元兜底 → 默认到总页数,会把 Workbook/词汇表/附录/封底算进来。
            # 裁剪到首个后置附录页之前(找不到则保持总页数,留人工调整)。
            bm = _backmatter_start(pages, start)
            end = (bm - 1) if bm else total
        segments.append({
            "unit_no": unit_no,
            "start_page": start,
            "end_page": max(start, end),
            "detected_title": title or None,
        })
    return segments


# ─── 文本拼接 ─────────────────────────────────────────────────────────────────

async def ocr_pages_to_sidecar(file_id: str, *, resolution: int = 130,
                               concurrency: int = 6, on_progress=None) -> int:
    """扫描件 PDF → 逐页渲染图片 → 豆包视觉 OCR → 原样文字,存 {file_id}.ocr.json。

    存好后 extract_pages 会自动改用 OCR 文字,单元检测/生成全链路透明。返回页数。
    """
    import asyncio
    import base64
    import io
    import json

    import pdfplumber
    from app.services.doubao_vision_service import recognize_page_text

    path = str(UPLOAD_DIR / f"{file_id}.pdf")

    def _render_all() -> list[str]:
        urls: list[str] = []
        with pdfplumber.open(path) as pdf:
            for pg in pdf.pages:
                pil = pg.to_image(resolution=resolution).original.convert("RGB")
                buf = io.BytesIO()
                pil.save(buf, format="JPEG", quality=80)
                urls.append("data:image/jpeg;base64," + base64.b64encode(buf.getvalue()).decode())
        return urls

    urls = await asyncio.to_thread(_render_all)
    n = len(urls)
    results = [""] * n
    done = 0
    sem = asyncio.Semaphore(concurrency)

    async def _one(i: int) -> None:
        nonlocal done
        async with sem:
            results[i] = (await recognize_page_text(urls[i]) or "").strip()
            done += 1
            if on_progress:
                on_progress(done, n)

    await asyncio.gather(*(_one(i) for i in range(n)))
    _ocr_sidecar(file_id).write_text(json.dumps(results, ensure_ascii=False), encoding="utf-8")
    return n


def split_unit_pdf(file_id: str, start_page: int, end_page: int) -> bytes:
    """从原始 PDF 抽取 [start_page, end_page](1-based 含首尾)生成单元独立 PDF bytes。

    文字版/扫描版都按原始页拆,得到的是"原版单元 PDF"。需 PyPDF2。
    """
    import io

    from PyPDF2 import PdfReader, PdfWriter
    path = UPLOAD_DIR / f"{file_id}.pdf"
    if not path.exists():
        raise FileNotFoundError(f"PDF not found: {file_id}")
    reader = PdfReader(str(path), strict=False)
    writer = PdfWriter()
    n = len(reader.pages)
    dropped: list[int] = []
    for i in range(max(0, start_page - 1), min(n, end_page)):
        try:
            writer.add_page(reader.pages[i])
        except Exception as exc:  # noqa: BLE001
            # 个别页有悬空对象引用（源 PDF 损坏 + PyPDF2 clone 的已知 bug，
            # 抛 AssertionError），跳过坏页保住整份单元 PDF，而非整份失败。
            dropped.append(i + 1)
            _log.warning("split_unit_pdf %s 第 %d 页无法克隆，已跳过：%s", file_id, i + 1, exc)
    if dropped:
        _log.warning("split_unit_pdf %s 跳过 %d 页：%s", file_id, len(dropped), dropped)
    if len(writer.pages) == 0:
        raise RuntimeError(f"split_unit_pdf 无可用页（{start_page}-{end_page}）")
    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


async def upload_pdf_to_cos(pdf_bytes: bytes, key: str) -> str | None:
    """上传 PDF bytes 到 COS(public-read),返回直链。COS 未配(dev)→ 返回 None。"""
    import asyncio

    from app.core.config import settings
    if settings.cos_secret_key.startswith("placeholder"):
        return None
    from app.services.vocab_media_provider import _get_cos_client

    def _put() -> None:
        _get_cos_client().put_object(
            Bucket=settings.cos_bucket, Key=key, Body=pdf_bytes,
            ContentType="application/pdf", ACL="public-read")

    await asyncio.to_thread(_put)
    return f"{settings.cos_base_url}/{key}"


def get_unit_text(file_id: str, start_page: int, end_page: int) -> str:
    """返回指定页范围（1-based 含首尾）的拼接文本，用于 AI 生成上下文。"""
    pages = extract_pages(file_id)
    selected = pages[start_page - 1: end_page]
    return "\n\n---\n\n".join(p for p in selected if p)


def get_page_previews(file_id: str) -> list[dict]:
    """返回每页摘要（前 200 字），供人工选区使用。"""
    pages = extract_pages(file_id)
    return [
        {"page_no": i + 1, "text_snippet": p[:200]}
        for i, p in enumerate(pages)
    ]
